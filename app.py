import streamlit as st
import requests
import json
import pandas as pd
from docx import Document
import io
import time
import concurrent.futures

# --- PAGE CONFIG ---
st.set_page_config(
    page_title="JEM English Subtitle Generator",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- SESSION STATE INITIALIZATION ---
if 'result' not in st.session_state:
    st.session_state['result'] = None
if 'confirm_clear' not in st.session_state:
    st.session_state['confirm_clear'] = False
if 'input_text' not in st.session_state:
    st.session_state['input_text'] = ""

# --- CALLBACKS ---
def on_text_change():
    """Clear previous results immediately when text changes."""
    st.session_state['result'] = None
    st.session_state['confirm_clear'] = False
    # Sync text area content to session state
    st.session_state['input_text'] = st.session_state.input_area

def request_clear():
    """Trigger the confirmation dialog."""
    st.session_state['confirm_clear'] = True

def confirm_clear_action():
    """Actually clear the data."""
    st.session_state['result'] = None
    st.session_state['input_text'] = ""
    st.session_state['confirm_clear'] = False
    st.rerun()

def cancel_clear():
    """Cancel the clear request."""
    st.session_state['confirm_clear'] = False

def handle_file_upload():
    """Read uploaded file (TXT or DOCX) and populate text area."""
    uploaded_file = st.session_state.uploaded_file
    if uploaded_file is not None:
        try:
            # Handle .docx files
            if uploaded_file.name.lower().endswith('.docx'):
                doc = Document(uploaded_file)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                file_content = '\n'.join(full_text)
            
            # Handle .txt files
            else:
                stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
                file_content = stringio.read()
            
            # Update state
            st.session_state['input_text'] = file_content
            st.session_state['result'] = None # Reset results
        except Exception as e:
            st.error(f"Error reading file: {e}")

# --- CUSTOM CSS ---
st.markdown("""
<style>
    /* IMPORT HEBREW FONTS */
    @import url('https://fonts.googleapis.com/css2?family=Frank+Ruhl+Libre:wght@400;700&family=Alef:wght@400;700&display=swap');

    /* MAIN BACKGROUND */
    .stApp {
        background-color: #FDFBF7;
        color: #4A3B32;
    }

    /* SIDEBAR BACKGROUND */
    section[data-testid="stSidebar"] {
        background-color: #F3F0E6;
        border-right: 1px solid #E0DACC;
    }

    /* HEADERS */
    h1, h2, h3, h4, .stMarkdown {
        color: #4A3B32 !important;
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
    }
    
    /* PRIMARY BUTTON (TRANSLATE) */
    div.stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #A67C52 0%, #8B5A2B 100%);
        color: white !important;
        border: none;
        border-radius: 8px;
        padding: 0.6rem 1.5rem;
        font-weight: 600;
        box-shadow: 0 4px 6px rgba(139, 90, 43, 0.2);
        transition: all 0.3s ease;
    }
    
    div.stButton > button[kind="primary"]:hover {
        background: linear-gradient(135deg, #8B5A2B 0%, #6F4E37 100%);
        box-shadow: 0 6px 8px rgba(139, 90, 43, 0.3);
        transform: translateY(-1px);
    }
    
    /* SECONDARY BUTTON (CLEAR) */
    div.stButton > button[kind="secondary"] {
        background-color: #FFFFFF !important;
        border: 2px solid #8B5A2B !important;
        color: #8B5A2B !important;
        font-weight: bold !important;
        border-radius: 8px;
        transition: all 0.3s ease;
    }
    
    div.stButton > button[kind="secondary"]:hover {
        background-color: #F3F0E6 !important;
        color: #6F4E37 !important;
        border-color: #6F4E37 !important;
    }

    /* INPUT TEXT AREA */
    .stTextArea textarea {
        background-color: #FFFFFF;
        border: 1px solid #D7D0C0;
        border-radius: 8px;
        color: #333333;
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.02);
    }
    
    /* FILE UPLOADER STYLE */
    [data-testid='stFileUploader'] {
        margin-bottom: 10px;
    }

    /* --- LOADING STEPS STYLING --- */
    .step-box {
        background-color: #FFF;
        border: 1px solid #E0DACC;
        border-radius: 8px;
        padding: 20px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        margin-bottom: 20px;
    }
    .step-item {
        font-family: 'Helvetica Neue', sans-serif;
        font-size: 1.05em;
        padding: 8px 0;
        color: #888;
        display: flex;
        align-items: center;
        transition: color 0.3s ease;
    }
    .step-item.active {
        color: #8B5A2B;
        font-weight: bold;
    }
    .step-item.done {
        color: #2E7D32; /* Green */
        font-weight: bold;
    }
    .step-icon {
        margin-right: 12px;
        width: 20px;
        display: inline-block;
        text-align: center;
    }

    /* --- TABLE STYLING --- */
    .results-table {
        width: 100%;
        border-collapse: collapse;
        margin-top: 20px;
        background-color: #FFF;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        border: 1px solid #E0DACC;
    }
    .results-table th {
        background-color: #EBE5D5;
        color: #4A3B32;
        padding: 12px 15px;
        text-align: left;
        font-weight: bold;
        border-bottom: 2px solid #D7D0C0;
        font-size: 0.95em;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .results-table td {
        padding: 12px 15px;
        border-bottom: 1px solid #F0EAE0;
        vertical-align: top; 
        color: #333;
    }
    .results-table tr:last-child td {
        border-bottom: none;
    }
    .results-table tr:hover {
        background-color: #FAF8F2;
    }
    .id-col { width: 50px; color: #8B5A2B; font-weight: bold; font-size: 0.85em; text-align: center; }
    .yiddish-col { font-family: 'Frank Ruhl Libre', 'Alef', serif; font-size: 1.3em; direction: rtl; text-align: right; color: #222; width: 45%; line-height: 1.5; }
    .english-col { font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; font-size: 1.05em; line-height: 1.5; width: 45%; color: #111; }
</style>
""", unsafe_allow_html=True)

# --- SIDEBAR (SETTINGS) ---
with st.sidebar:
    st.markdown("### Settings")
    api_key = st.text_input("Enter Google API Key", type="password")
    
    default_prompt = """
# Role
You are a master subtitler adapting the Lubavitcher Rebbe’s Sichos. Your goal is to produce **narrative, high-impact English** that captures the speaker's voice while adhering to strict video-subtitle standards.

# MODULE A: PERSPECTIVE (The "No Narrator" Rule)
* **CRITICAL:** These are subtitles for the person on screen.
* **FORBIDDEN:** Never write "The Rebbe explains," "He emphasizes," "The speaker continues," or "They teach us."
* **ACTION:** Translate ONLY what is said in the first person (I/We).
    * *Bad:* "The Rebbe explains that the Alter Rebbe was released..."
    * *Good:* "Although **the Alter Rebbe** was released..."

# MODULE B: FIDELITY & AGGRESSIVE SEGMENTATION
* **The "Zero-Loss" Rule:** Do not summarize. Every distinct thought in the Yiddish must have a corresponding English phrase.
* **The "Short-Burst" Rule:** Do not try to fit a long Yiddish sentence into one subtitle. **Break long complex sentences into 2, 3, or even 4 separate, short subtitle events.**
    * *Logic:* Better to have 3 short, readable subtitles than 1 long, crowded one.

# MODULE C: NARRATIVE VOICE & THEOLOGY
### 1. VOICE ATTRIBUTION (Internal Dialogue)
* **Action:** Insert tags to describe internal thought processes of *characters in the story* (not the speaker).
    * *Input:* "If the other person..." -> *Output:* "**Moses reasoned**, 'If another person...'"

### 2. PHENOMENON OVER LABEL
* **Action:** Describe the effect/meaning, not the technical label.
    * *Input:* "Nimna Hanimnaos" -> *Output:* "**God, who is Infinite, and therefore contains the finite.**"

### 3. THE "QUOTE CONTEXT" RULE
* **Liturgy:** Translate objects of study literally ("**Hear O Israel...**").
* **Prooftext:** Translate the point of the quote ("**Man was born to toil**").

# MODULE D: CULTURAL & LINGUISTIC TRANSLATION
### 4. CONCEPT OVER ETYMOLOGY
* **Action:** Translate the *implication*, not the literal word.
    * *Input:* "Chozer L'buryo" -> "**Returns to full health**" (NOT "Returns to his wholeness").
    * *Input:* "Adam" -> "**Created in God's image.**"

### 5. MECHANICS VS. MEANING
* **Action:** If text uses mechanics (Gematria/Letters) to explain a concept, translate the **concept**.

### 6. RELATIONAL TITLES
* **Action:** *Der Rebbe (Nishmaso Eden)* -> "**My father-in-law, the Rebbe.**"

# MODULE E: VISUAL STRUCTURE & RHYTHM
* **Length:** Max 42 characters per line.
* **Balance:** If using 2 lines, keep them roughly equal in length.
* **Split:** Use the tilde symbol `~` to indicate a visual line break inside a single subtitle row.

# MODULE F: SYNTAX (The "Manual" Rules)
* **Active Voice:** "It is believed by many" -> "**Many believe**"
* **No Double Negatives:** "Not only will they not disturb" -> "**The government will not disturb; / on the contrary, it will help.**"

# FEW-SHOT EXAMPLES (Correct Style)

### Example 1: Removing "The Rebbe Explains" & Segmentation
*Input:*
וואָרום אף על פי וואָס דער שחרור פון דעם אַלטן רבי'ן איז געווען ביום י"ט כסלו, איז דאָך געווען כמה סיבות וואָס דערפאַר האָט זיך פאַרהאַלטן
*Output:*
001 | וואָרום אף על פי וואָס דער שחרור... | And yes, **the Alter Rebbe**~was technically released on the 19th.
002 | איז דאָך געווען כמה סיבות וואָס דערפאַר האָט זיך פאַרהאַלטן | However, due to various reasons,~he was detained.

### Example 2: Idiomatic Translation (No "Wholeness")
*Input:*
נאָר אויך ער דאַרף צוואוואַרטן ביז "חוזר לבוריו" – ער ווערט אינגאַנצן געזונט
*Output:*
010 | נאָר אויך ער דאַרף צוואוואַרטן ביז "חוזר לבוריו" | One waits until he~"returns to full health."
011 | ער ווערט אינגאַנצן געזונט | Only then is the recovery complete.

### Example 3: Voice Attribution (Internal Character)
*Input:*
משה האט געטראכט אז אויב יענער טוט אזוי...
*Output:*
020 | משה האט געטראכט אז אויב יענער טוט אזוי... | **Moses reasoned:**~"If that person acts this way..."

# TECHNICAL OUTPUT FORMAT
Provide the output as a simple list with 3 columns separated by pipes (|).
Do NOT use Markdown Table syntax. Just raw lines.

Format:
ID | Yiddish Snippet | English Subtitle
    """
    
    with st.expander("Edit System Prompt"):
        system_prompt = st.text_area("Prompt", value=default_prompt, height=400)

# --- API FUNCTION (SYNC) ---
def call_api(model_name, api_key, full_prompt):
    url = f"https://generativelanguage.googleapis.com/v1beta/{model_name}:generateContent?key={api_key}"
    headers = {'Content-Type': 'application/json'}
    data = {
        "contents": [{"parts": [{"text": full_prompt}]}],
        "safetySettings": [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
        ]
    }
    
    try:
        response = requests.post(url, headers=headers, data=json.dumps(data))
        if response.status_code == 200:
            result_json = response.json()
            try:
                text = result_json['candidates'][0]['content']['parts'][0]['text']
                return True, text
            except KeyError:
                return False, f"Parsed JSON but found no text: {result_json}"
        elif response.status_code == 404:
            return False, "NOT_FOUND"
        else:
            return False, f"Error {response.status_code}: {response.text}"
    except Exception as e:
        return False, str(e)

# --- RENDERER FOR LOADING ANIMATION ---
def render_steps(steps, current_idx):
    html_content = '<div class="step-box">'
    html_content += '<div style="margin-bottom:10px; font-weight:bold; color:#4A3B32;">Bringing the Rebbe to the English speaking world...</div>'
    
    for i, step_text in enumerate(steps):
        if i < current_idx:
            # Completed
            icon = "✅"
            cls = "step-item done"
        elif i == current_idx:
            # Active
            icon = "🔄"
            cls = "step-item active"
        else:
            # Pending
            icon = "⚪"
            cls = "step-item"
        
        html_content += f'<div class="{cls}"><span class="step-icon">{icon}</span>{step_text}</div>'
    
    html_content += '</div>'
    return html_content

# --- HEADER (LOGO RIGHT) ---
header_col1, header_col2 = st.columns([5, 1])
with header_col1:
    st.title("JEM English Subtitle Generator For Sichos")
with header_col2:
    st.image("https://cdn.prod.website-files.com/67dd48176b6e7b21cf6cc9bc/67ecd01b1392af3e91a89ee2_d68e8bf608547d781d6eaf13a23203df_jem%20ai%20hero.svg", use_column_width=True)

st.markdown("---")

# --- MAIN LAYOUT ---
col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.markdown("### Input Transcript")
    
    # File Uploader
    st.file_uploader(
        "Upload File (.txt or .docx)", 
        type=["txt", "docx"], 
        key="uploaded_file", 
        on_change=handle_file_upload,
        label_visibility="collapsed"
    )

    # Text Area
    yiddish_text = st.text_area(
        "Paste Yiddish text here...", 
        height=600, 
        key="input_area",
        value=st.session_state['input_text'],
        on_change=on_text_change
    )

with col2:
    st.markdown("### Generated Subtitles")
    
    # Buttons
    b_col1, b_col2 = st.columns([3, 2])
    with b_col1:
        translate_btn = st.button("TRANSLATE", type="primary", use_container_width=True)
    with b_col2:
        clear_btn = st.button("CLEAR", type="secondary", use_container_width=True, on_click=request_clear)

    # Confirmation
    if st.session_state.get('confirm_clear'):
        st.warning("⚠️ **Are you sure?** Unsaved translations will be lost.")
        conf_col1, conf_col2 = st.columns([1, 1])
        with conf_col1:
            st.button("Yes, Clear Everything", type="primary", use_container_width=True, on_click=confirm_clear_action)
        with conf_col2:
            st.button("No, Cancel", type="secondary", use_container_width=True, on_click=cancel_clear)

    # Logic Container
    result_container = st.container()

    if translate_btn and not st.session_state.get('confirm_clear'):
        if not api_key:
            st.error("Please enter your API Key in the sidebar.")
        elif not yiddish_text:
            st.warning("Please paste text to translate.")
        else:
            combined_prompt = f"{system_prompt}\n\n---\n\nTASK: Translate this text:\n{yiddish_text}"
            
            # --- THE PACED LOADING SIMULATION ---
            loading_placeholder = st.empty()
            
            # Steps to display
            steps = [
                "Analyzing First-Person Perspective...",
                "Applying Segmentation Rules...",
                "Verifying Narrative Voice & Context...",
                "Checking Cultural Nuances...",
                "Optimizing Visual Rhythm...",
                "Finalizing Syntax & Grammar..."
            ]
            
            # Start API in Thread
            executor = concurrent.futures.ThreadPoolExecutor()
            future = executor.submit(call_api, "models/gemini-2.5-flash", api_key, combined_prompt)
            
            # Animation Loop
            total_steps = len(steps)
            current_step = 0
            
            while current_step < total_steps:
                loading_placeholder.markdown(render_steps(steps, current_step), unsafe_allow_html=True)
                time.sleep(1.2)
                current_step += 1
            
            while not future.done():
                 loading_placeholder.markdown(render_steps(steps, total_steps - 1), unsafe_allow_html=True)
                 time.sleep(0.5)

            # Get Result
            success, result = future.result()
            
            # If not found, try backup
            if not success and result == "NOT_FOUND":
                 loading_placeholder.info("Trying backup model...")
                 success, result = call_api("models/gemini-1.5-flash", api_key, combined_prompt)

            # Clear Loader
            loading_placeholder.empty()

            if success:
                st.session_state['result'] = result
            else:
                st.error(f"❌ Failed: {result}")

    # --- RESULTS DISPLAY ---
    if st.session_state.get('result') and not st.session_state.get('confirm_clear'):
        raw_text = st.session_state['result']
        
        # Parse Data
        data = []
        for line in raw_text.split('\n'):
            if "|" in line and "ID |" not in line and "---" not in line:
                parts = line.split('|')
                if len(parts) >= 3:
                    data.append({
                        "id": parts[0].strip(),
                        "yiddish": parts[1].strip(),
                        "english_clean": parts[2].strip().replace("~", " "),
                        "english_raw": parts[2].strip().replace("~", "\n")
                    })
        
        with result_container:
            if data:
                # Build HTML Table
                table_html = """<table class="results-table">
<thead>
<tr>
<th style="width:50px;">ID</th>
<th style="text-align:right;">Yiddish Source</th>
<th>English Subtitle</th>
</tr>
</thead>
<tbody>"""
                for row in data:
                    table_html += f"""<tr>
<td class="id-col">{row['id']}</td>
<td class="yiddish-col">{row['yiddish']}</td>
<td class="english-col">{row['english_clean']}</td>
</tr>"""
                table_html += "</tbody></table>"
                
                st.markdown(table_html, unsafe_allow_html=True)
                
                # DOCX Export
                df = pd.DataFrame(data)
                doc = Document()
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                for idx, row in df.iterrows():
                    cells = table.add_row().cells
                    cells[0].text = row['id']
                    cells[1].text = row['yiddish']
                    cells[2].text = row['english_raw']
                
                bio = io.BytesIO()
                doc.save(bio)
                st.markdown("<br>", unsafe_allow_html=True)
                st.download_button("Download DOCX", bio.getvalue(), "translation.docx", use_container_width=True)
            
            with st.expander("View Raw Output"):
                st.text(raw_text)
